from __future__ import annotations

import json
import math
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import networkx as nx
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DIAGRAM_DIR = ROOT / "diagrams"
OUTPUT = ROOT / "Cafe_Food_Ordering_System_Project_Report.docx"
VALIDATION = ROOT / "diagram_validation.json"
DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)

PROJECT = "Cafe: Online Food Ordering and Restaurant Management System"
STUDENT = "Sumit Kumar"
INSTITUTION = "[Enter College / University Name]"
DEPARTMENT = "[Enter Department Name]"
ROLL = "[Enter Roll Number]"
SUPERVISOR = "[Enter Project Supervisor Name]"
SESSION = "2025-2026"


def diagram_specs():
    return [
        {
            "id": "01_system_context",
            "title": "System Context Diagram",
            "nodes": {
                "Customer": (-3, 1.8), "Admin": (-3, 0.6), "Kitchen": (-3, -0.6),
                "Delivery": (-3, -1.8), "Cafe System": (0, 0), "MariaDB": (3, 0),
            },
            "edges": [
                ("Customer", "Cafe System", "browse, order, review"),
                ("Admin", "Cafe System", "manage and report"),
                ("Kitchen", "Cafe System", "prepare and update"),
                ("Delivery", "Cafe System", "deliver and confirm"),
                ("Cafe System", "MariaDB", "validated SQL"),
                ("MariaDB", "Cafe System", "stored records"),
            ],
        },
        {
            "id": "02_layered_architecture",
            "title": "Layered Application Architecture",
            "nodes": {
                "Web Browser": (-3, 1.2), "HTML/CSS/JS": (-1.5, 1.2),
                "PHP Pages": (0, 1.2), "Actions Controller": (0, 0),
                "Auth / CSRF": (-1.8, 0), "Business Helpers": (1.8, 0),
                "PDO Data Access": (0, -1.2), "MariaDB": (0, -2.4),
            },
            "edges": [
                ("Web Browser", "HTML/CSS/JS", "request / interaction"),
                ("HTML/CSS/JS", "PHP Pages", "GET render"),
                ("PHP Pages", "Web Browser", "HTML response"),
                ("HTML/CSS/JS", "Actions Controller", "POST form"),
                ("Actions Controller", "Auth / CSRF", "authorize"),
                ("Actions Controller", "Business Helpers", "calculate"),
                ("Actions Controller", "PDO Data Access", "prepared query"),
                ("PHP Pages", "PDO Data Access", "read models"),
                ("PDO Data Access", "MariaDB", "SQL"),
                ("MariaDB", "PDO Data Access", "rows"),
            ],
        },
        {
            "id": "03_customer_use_cases",
            "title": "Customer Use-Case Diagram",
            "nodes": {
                "Customer": (-3, 0), "Register / Login": (-0.7, 2),
                "Browse Menu": (-0.7, 1), "Manage Cart": (-0.7, 0),
                "Checkout": (-0.7, -1), "Track Order": (1.8, 1),
                "Favorites": (1.8, 0), "Review Food": (1.8, -1),
            },
            "edges": [
                ("Customer", "Register / Login", "starts"),
                ("Customer", "Browse Menu", "uses"),
                ("Customer", "Manage Cart", "uses"),
                ("Customer", "Checkout", "places order"),
                ("Customer", "Track Order", "monitors"),
                ("Customer", "Favorites", "saves"),
                ("Customer", "Review Food", "rates delivered item"),
                ("Browse Menu", "Manage Cart", "add item"),
                ("Manage Cart", "Checkout", "proceed"),
                ("Checkout", "Track Order", "order created"),
            ],
        },
        {
            "id": "04_staff_use_cases",
            "title": "Role-Based Staff Use Cases",
            "nodes": {
                "Administrator": (-3, 1.5), "Kitchen Staff": (-3, 0),
                "Delivery Staff": (-3, -1.5), "Menu CRUD": (0, 2),
                "Order Control": (0, 1), "Reports / Export": (0, 0),
                "Prepare Order": (2.8, 0.8), "Mark Ready": (2.8, 0),
                "Start Delivery": (0, -1.2), "Mark Delivered": (2.8, -1.2),
            },
            "edges": [
                ("Administrator", "Menu CRUD", "manages"),
                ("Administrator", "Order Control", "controls"),
                ("Administrator", "Reports / Export", "analyzes"),
                ("Kitchen Staff", "Prepare Order", "starts"),
                ("Prepare Order", "Mark Ready", "completes"),
                ("Delivery Staff", "Start Delivery", "accepts"),
                ("Start Delivery", "Mark Delivered", "completes"),
                ("Order Control", "Prepare Order", "confirms order"),
                ("Mark Ready", "Start Delivery", "handover"),
            ],
        },
        {
            "id": "05_dfd_level_0",
            "title": "Data Flow Diagram - Level 0",
            "nodes": {
                "Customer": (-3.2, 1.2), "Staff Roles": (-3.2, -1.2),
                "1.0 Authentication": (-0.8, 1.5), "2.0 Ordering": (-0.8, 0),
                "3.0 Management": (-0.8, -1.5), "User Store": (2.5, 1.5),
                "Order Store": (2.5, 0), "Menu Store": (2.5, -1.5),
            },
            "edges": [
                ("Customer", "1.0 Authentication", "credentials"),
                ("1.0 Authentication", "User Store", "user query"),
                ("User Store", "1.0 Authentication", "account"),
                ("Customer", "2.0 Ordering", "cart and checkout"),
                ("2.0 Ordering", "Order Store", "new order"),
                ("Menu Store", "2.0 Ordering", "price and stock"),
                ("Staff Roles", "3.0 Management", "CRUD / status"),
                ("3.0 Management", "Menu Store", "menu changes"),
                ("3.0 Management", "Order Store", "status changes"),
                ("Order Store", "Customer", "tracking status"),
            ],
        },
        {
            "id": "06_order_dfd_level_1",
            "title": "Order Processing DFD - Level 1",
            "nodes": {
                "Cart Session": (-3, 1.5), "Validate Items": (-1, 1.5),
                "Validate Coupon": (1, 1.5), "Calculate Totals": (3, 1.5),
                "Create Order": (3, 0), "Create Items": (1, 0),
                "Update Stock": (-1, 0), "Commit Transaction": (-3, 0),
                "Notify Customer": (0, -1.5),
            },
            "edges": [
                ("Cart Session", "Validate Items", "quantity and IDs"),
                ("Validate Items", "Validate Coupon", "valid cart"),
                ("Validate Coupon", "Calculate Totals", "discount"),
                ("Calculate Totals", "Create Order", "financial values"),
                ("Create Order", "Create Items", "order ID"),
                ("Create Items", "Update Stock", "quantities"),
                ("Update Stock", "Commit Transaction", "all updates pass"),
                ("Commit Transaction", "Notify Customer", "committed order"),
            ],
        },
        {
            "id": "07_er_diagram",
            "title": "Entity Relationship Diagram",
            "nodes": {
                "users": (-3, 2), "addresses": (-1, 3), "orders": (-1, 1.2),
                "order_items": (1.2, 1.2), "menu_items": (3.3, 1.2),
                "categories": (3.3, 3), "reviews": (1.2, -0.5),
                "favorites": (3.3, -0.5), "notifications": (-3, 0),
                "review_helpful": (1.2, -2.1), "coupons": (-1, -2.1),
                "settings": (3.3, -2.1), "activity_logs": (-3, -2.1),
            },
            "edges": [
                ("addresses", "users", "user_id FK"),
                ("orders", "users", "user_id FK"),
                ("order_items", "orders", "order_id FK"),
                ("order_items", "menu_items", "menu_item_id FK"),
                ("menu_items", "categories", "category_id FK"),
                ("reviews", "users", "user_id FK"),
                ("reviews", "orders", "order_id FK"),
                ("reviews", "menu_items", "menu_item_id FK"),
                ("favorites", "users", "user_id FK"),
                ("favorites", "menu_items", "menu_item_id FK"),
                ("notifications", "users", "user_id FK"),
                ("review_helpful", "reviews", "review_id FK"),
                ("review_helpful", "users", "user_id FK"),
                ("activity_logs", "users", "user_id FK"),
            ],
        },
        {
            "id": "08_login_sequence",
            "title": "Login Sequence",
            "nodes": {
                "User": (-3, 0), "Login Form": (-1.5, 0), "actions.php": (0, 0),
                "PDO": (1.5, 0), "users table": (3, 0), "Session": (0, -1.5),
                "Role Dashboard": (2.2, -1.5),
            },
            "edges": [
                ("User", "Login Form", "enter credentials"),
                ("Login Form", "actions.php", "POST + CSRF"),
                ("actions.php", "PDO", "prepared SELECT"),
                ("PDO", "users table", "email lookup"),
                ("users table", "PDO", "account row"),
                ("PDO", "actions.php", "result"),
                ("actions.php", "Session", "regenerate + store user"),
                ("Session", "Role Dashboard", "authorized redirect"),
            ],
        },
        {
            "id": "09_checkout_sequence",
            "title": "Checkout and Order Creation Sequence",
            "nodes": {
                "Customer": (-3, 1), "Checkout Form": (-1.5, 1),
                "Order Handler": (0, 1), "Menu Validation": (1.5, 1),
                "DB Transaction": (3, 1), "Orders": (3, -0.5),
                "Order Items": (1.5, -0.5), "Stock": (0, -0.5),
                "Success Page": (-1.5, -0.5),
            },
            "edges": [
                ("Customer", "Checkout Form", "submit details"),
                ("Checkout Form", "Order Handler", "POST + token"),
                ("Order Handler", "Menu Validation", "re-fetch prices"),
                ("Menu Validation", "DB Transaction", "validated cart"),
                ("DB Transaction", "Orders", "INSERT order"),
                ("Orders", "Order Items", "new order ID"),
                ("Order Items", "Stock", "decrement quantity"),
                ("Stock", "DB Transaction", "commit"),
                ("DB Transaction", "Success Page", "redirect with order ID"),
            ],
        },
        {
            "id": "10_order_state",
            "title": "Order Status State Machine",
            "nodes": {
                "Pending": (-3, 0), "Confirmed": (-1.8, 0), "Preparing": (-0.4, 0),
                "Ready": (1, 0), "Out for Delivery": (2.4, 0),
                "Delivered": (4.1, 0), "Cancelled": (-0.4, -1.5),
            },
            "edges": [
                ("Pending", "Confirmed", "admin / staff"),
                ("Confirmed", "Preparing", "kitchen"),
                ("Preparing", "Ready", "kitchen"),
                ("Ready", "Out for Delivery", "rider assigned"),
                ("Out for Delivery", "Delivered", "delivery staff"),
                ("Pending", "Cancelled", "authorized cancellation"),
                ("Confirmed", "Cancelled", "authorized cancellation"),
            ],
        },
        {
            "id": "11_deployment",
            "title": "Local Deployment Diagram",
            "nodes": {
                "User Browser": (-3, 1), "Apache Server": (-1, 1),
                "PHP 8 Runtime": (1, 1), "MariaDB": (3, 1),
                "Project Files": (1, -0.8), "Uploaded Images": (3, -0.8),
                "phpMyAdmin": (-1, -0.8),
            },
            "edges": [
                ("User Browser", "Apache Server", "HTTP localhost"),
                ("Apache Server", "PHP 8 Runtime", "execute .php"),
                ("PHP 8 Runtime", "MariaDB", "PDO connection"),
                ("Project Files", "PHP 8 Runtime", "include / render"),
                ("PHP 8 Runtime", "Uploaded Images", "store path"),
                ("phpMyAdmin", "MariaDB", "import SQL"),
            ],
        },
        {
            "id": "12_module_dependency",
            "title": "Module Dependency Diagram",
            "nodes": {
                "config.php": (0, 2), "functions.php": (0, 1),
                "index.php": (-2.5, 0), "dashboard.php": (-0.8, 0),
                "actions.php": (0.9, 0), "export.php": (2.5, 0),
                "style.css": (-2.5, -1.5), "app.js": (-0.8, -1.5),
                "food_ordering.sql": (1.7, -1.5),
            },
            "edges": [
                ("functions.php", "config.php", "requires"),
                ("index.php", "functions.php", "requires"),
                ("dashboard.php", "functions.php", "requires"),
                ("actions.php", "functions.php", "requires"),
                ("export.php", "functions.php", "requires"),
                ("index.php", "style.css", "loads"),
                ("index.php", "app.js", "loads"),
                ("dashboard.php", "style.css", "loads"),
                ("dashboard.php", "app.js", "loads"),
                ("config.php", "food_ordering.sql", "database name"),
            ],
        },
        {
            "id": "13_security_flow",
            "title": "Request Security Validation Flow",
            "nodes": {
                "Incoming POST": (-3, 1), "Method Check": (-1.5, 1),
                "CSRF Verify": (0, 1), "Session Auth": (1.5, 1),
                "Role Guard": (3, 1), "Input Validation": (3, -0.5),
                "Prepared Statement": (1.5, -0.5), "Transaction": (0, -0.5),
                "Escaped Response": (-1.5, -0.5), "Reject 4xx": (0, -2),
            },
            "edges": [
                ("Incoming POST", "Method Check", "request"),
                ("Method Check", "CSRF Verify", "POST"),
                ("CSRF Verify", "Session Auth", "valid token"),
                ("Session Auth", "Role Guard", "logged in"),
                ("Role Guard", "Input Validation", "permitted"),
                ("Input Validation", "Prepared Statement", "clean values"),
                ("Prepared Statement", "Transaction", "safe SQL"),
                ("Transaction", "Escaped Response", "commit / redirect"),
                ("Method Check", "Reject 4xx", "wrong method"),
                ("CSRF Verify", "Reject 4xx", "invalid token"),
                ("Role Guard", "Reject 4xx", "forbidden"),
            ],
        },
        {
            "id": "14_navigation",
            "title": "Application Navigation Structure",
            "nodes": {
                "Public Home": (0, 2.2), "Login / Register": (-2.8, 1),
                "Menu / Item": (-0.9, 1), "Customer Dashboard": (1, 1),
                "Role Dashboard": (2.9, 1), "Cart / Checkout": (-0.9, -0.5),
                "Orders / Tracking": (1, -0.5), "Admin Modules": (2.9, -0.5),
                "Kitchen Queue": (2, -2), "Delivery Hub": (3.8, -2),
            },
            "edges": [
                ("Public Home", "Login / Register", "account"),
                ("Public Home", "Menu / Item", "browse"),
                ("Login / Register", "Customer Dashboard", "customer"),
                ("Login / Register", "Role Dashboard", "staff"),
                ("Menu / Item", "Cart / Checkout", "add / buy"),
                ("Cart / Checkout", "Orders / Tracking", "placed"),
                ("Customer Dashboard", "Orders / Tracking", "view"),
                ("Role Dashboard", "Admin Modules", "admin"),
                ("Role Dashboard", "Kitchen Queue", "kitchen"),
                ("Role Dashboard", "Delivery Hub", "delivery"),
            ],
        },
        {
            "id": "15_test_pipeline",
            "title": "Verification and Test Pipeline",
            "nodes": {
                "Requirement": (-3, 0), "Test Case": (-1.8, 0),
                "HTTP Request": (-0.5, 0), "PHP Handler": (0.8, 0),
                "Database Check": (2.1, 0), "UI Response": (3.4, 0),
                "Pass / Fail Log": (0.8, -1.5), "Clean Test Data": (2.7, -1.5),
            },
            "edges": [
                ("Requirement", "Test Case", "derive"),
                ("Test Case", "HTTP Request", "execute"),
                ("HTTP Request", "PHP Handler", "route"),
                ("PHP Handler", "Database Check", "persist"),
                ("Database Check", "UI Response", "verify response"),
                ("UI Response", "Pass / Fail Log", "record"),
                ("Pass / Fail Log", "Clean Test Data", "after test"),
            ],
        },
    ]


def validate_and_render(spec):
    nodes = spec["nodes"]
    edges = spec["edges"]
    errors = []
    for source, target, label in edges:
        if source not in nodes:
            errors.append(f"Missing source node: {source}")
        if target not in nodes:
            errors.append(f"Missing target node: {target}")
        if source == target:
            errors.append(f"Self-loop is not permitted: {source}")
        if not label.strip():
            errors.append(f"Missing arrow label: {source} -> {target}")
    if errors:
        raise ValueError(f"{spec['id']}: " + "; ".join(errors))

    graph = nx.DiGraph()
    graph.add_nodes_from(nodes)
    graph.add_edges_from((a, b) for a, b, _ in edges)
    fig, ax = plt.subplots(figsize=(12.5, 7.1), dpi=180)
    ax.set_facecolor("#fbfaf7")
    pos = nodes
    nx.draw_networkx_nodes(
        graph, pos, node_size=3200, node_color="#fff4dc",
        edgecolors="#e29b1f", linewidths=1.6, ax=ax,
    )
    nx.draw_networkx_labels(
        graph, pos, font_size=8.2, font_weight="bold",
        font_color="#211d18", ax=ax,
    )
    nx.draw_networkx_edges(
        graph, pos, edge_color="#4d4a46", width=1.35,
        arrows=True, arrowstyle="-|>", arrowsize=18,
        connectionstyle="arc3,rad=0.04", min_source_margin=24,
        min_target_margin=24, ax=ax,
    )
    edge_labels = {(a, b): label for a, b, label in edges}
    nx.draw_networkx_edge_labels(
        graph, pos, edge_labels=edge_labels, font_size=6.3,
        rotate=False, bbox={"facecolor": "white", "alpha": 0.88, "edgecolor": "none", "pad": 1},
        ax=ax,
    )
    ax.set_title(spec["title"], fontsize=15, fontweight="bold", color="#2c241c", pad=14)
    xs = [point[0] for point in pos.values()]
    ys = [point[1] for point in pos.values()]
    ax.set_xlim(min(xs) - 1.2, max(xs) + 1.2)
    ax.set_ylim(min(ys) - 1.0, max(ys) + 1.0)
    ax.axis("off")
    fig.tight_layout()
    output = DIAGRAM_DIR / f"{spec['id']}.png"
    fig.savefig(output, bbox_inches="tight", facecolor="#fbfaf7")
    plt.close(fig)
    return {
        "id": spec["id"], "title": spec["title"], "node_count": len(nodes),
        "arrow_count": len(edges), "valid": True,
        "rule": "Every directed arrow source and target exists; arrowhead points from first node to second node.",
        "file": str(output.relative_to(ROOT)),
    }


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and choose Update Field if this table is not populated."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, separate, text, end])


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.72)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Title", 24, RGBColor(155, 92, 10)),
        ("Heading 1", 18, RGBColor(155, 92, 10)),
        ("Heading 2", 14, RGBColor(54, 44, 34)),
        ("Heading 3", 12, RGBColor(54, 44, 34)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = PROJECT
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = RGBColor(120, 115, 108)
        add_page_number(section.footer.paragraphs[0])


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(155, 92, 10)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(12)


def paragraph(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[idx], header, bold=True, color=(255, 255, 255), size=9)
        shade_cell(tbl.rows[0].cells[idx], "9B5C0A")
    for row_idx, row in enumerate(rows):
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=8.5)
            if row_idx % 2:
                shade_cell(cells[idx], "FFF8E8")
    if widths:
        for row in tbl.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return tbl


def figure(doc, filename, caption, number):
    doc.add_picture(str(filename), width=Inches(6.7))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Figure {number}: {caption}")
    r.bold = True
    r.font.size = Pt(9)


def new_report_page(doc, heading, chapter=None):
    if len(doc.paragraphs) > 0:
        doc.add_page_break()
    if chapter:
        p = doc.add_paragraph(chapter.upper())
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(150, 145, 138)
        p.runs[0].font.bold = True
    doc.add_heading(heading, level=1)


def add_standard_narrative(doc, subject, details, operation, outcomes):
    paragraph(doc, f"{subject} is an important part of the Cafe platform because it connects visible user interaction with reliable server-side processing. The implementation was designed for a local XAMPP environment while retaining patterns that can be migrated to a hosted Apache and MariaDB deployment.")
    paragraph(doc, f"In practical operation, {details} In this project, {operation}. Requests are checked at the server boundary, data is exchanged through prepared PDO statements, and the resulting state is rendered back into a responsive interface.")
    paragraph(doc, f"The resulting outcome is {outcomes}. The design also provides traceability for testing because each interface action maps to a named action handler, one or more database operations, and an observable response.")


def build_document(diagrams):
    doc = Document()
    configure_document(doc)

    # 1 - Cover
    title(doc, "PROJECT REPORT", "Submitted in partial fulfillment of the requirements for the award of the degree")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(25)
    r = p.add_run(PROJECT.upper())
    r.bold = True
    r.font.size = Pt(23)
    r.font.color.rgb = RGBColor(34, 29, 24)
    for text in [
        "", f"Submitted by\n{STUDENT}\nRoll Number: {ROLL}",
        f"\nUnder the Guidance of\n{SUPERVISOR}",
        f"\n{DEPARTMENT}\n{INSTITUTION}\nAcademic Session {SESSION}",
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(12)
            if STUDENT in run.text or INSTITUTION in run.text:
                run.bold = True

    # 2 - Certificate
    new_report_page(doc, "Certificate")
    paragraph(doc, f"This is to certify that the project report entitled “{PROJECT}” is a bona fide record of work carried out by {STUDENT}, Roll Number {ROLL}, under my supervision during the academic session {SESSION}.")
    paragraph(doc, "To the best of my knowledge, the work presented in this report has not been submitted, either in full or in part, for the award of any other degree or diploma. The system was reviewed with respect to its stated objectives, implementation, database design, testing evidence, and documentation.")
    table(doc, ["Project Guide", "Head of Department", "External Examiner"], [
        [f"{SUPERVISOR}\nSignature: __________", "Name: __________\nSignature: __________", "Name: __________\nSignature: __________"]
    ])
    paragraph(doc, "Place: ____________________")
    paragraph(doc, "Date: _____________________")

    # 3 - Declaration
    new_report_page(doc, "Student Declaration")
    paragraph(doc, f"I, {STUDENT}, declare that the report titled “{PROJECT}” is my original project work. All external ideas, software libraries, images, documentation sources, and standards referenced in this report have been acknowledged in the bibliography.")
    paragraph(doc, "I understand the principles of academic integrity and accept responsibility for ensuring that institution-specific details, screenshots, signatures, and deployment credentials are reviewed before final submission.")
    paragraph(doc, "\n\nStudent Signature: __________________________")
    paragraph(doc, f"Name: {STUDENT}")
    paragraph(doc, f"Roll Number: {ROLL}")

    # 4 - Acknowledgement
    new_report_page(doc, "Acknowledgement")
    paragraph(doc, f"I express sincere gratitude to {SUPERVISOR} for guidance, constructive feedback, and encouragement throughout the planning, development, testing, and documentation of this project.")
    paragraph(doc, f"I also thank the faculty members of {DEPARTMENT}, {INSTITUTION}, for providing the academic environment and technical foundation required to complete this work. Appreciation is extended to friends and family members who supported usability review and demonstration activities.")
    paragraph(doc, "Finally, I acknowledge the maintainers of PHP, MariaDB, Apache, Bootstrap, Bootstrap Icons, Python, and other open technologies used for development and documentation.")

    # 5 - Abstract
    new_report_page(doc, "Abstract")
    paragraph(doc, "Cafe is a browser-based food ordering and restaurant management system implemented using Core PHP, MariaDB/MySQL, HTML, CSS, JavaScript, Bootstrap, PDO, and XAMPP. The system combines a public delivery landing page, customer menu and checkout experience, visual order tracking, and role-specific operational dashboards.")
    paragraph(doc, "Customers can register, authenticate, browse and filter menu items, manage quantities, apply coupons, select delivery or pickup, place orders, save favorites, monitor status, and review delivered food. Administrators manage orders, menu items, food images, customers, staff roles, coupons, reviews, settings, reports, and CSV exports. Kitchen staff and delivery staff receive focused workflow screens.")
    paragraph(doc, "The project uses server-side validation, password hashing, session regeneration, CSRF tokens, output escaping, prepared SQL statements, guarded role transitions, transaction-based order placement, stock verification, and secure image validation. The result is an educational full-stack system that demonstrates practical database-backed web application engineering.")
    table(doc, ["Attribute", "Value"], [
        ["Application type", "Responsive full-stack web application"],
        ["Backend", "Core PHP with PDO"],
        ["Database", "MariaDB / MySQL"],
        ["Frontend", "HTML5, CSS3, JavaScript, Bootstrap 5"],
        ["Primary environment", "Apache and MariaDB supplied through XAMPP"],
        ["Supported roles", "Customer, Administrator, Kitchen, Delivery, Staff"],
    ])

    # 6 - TOC
    new_report_page(doc, "Table of Contents")
    add_toc(doc.add_paragraph())
    paragraph(doc, "The table of contents is generated from Word heading styles. Open the document in Microsoft Word, select the table, and choose Update Field if page numbers are not refreshed automatically.")

    # 7 - List of Figures
    new_report_page(doc, "List of Figures")
    rows = [[f"Figure {i}", item["title"], item["file"]] for i, item in enumerate(diagrams, 1)]
    table(doc, ["Number", "Figure Title", "Validated Asset"], rows)

    # 8 - List of Tables
    new_report_page(doc, "List of Tables")
    table(doc, ["Table", "Description"], [
        ["1", "Technology stack"],
        ["2", "Stakeholder requirements"],
        ["3", "Functional requirements"],
        ["4", "Non-functional requirements"],
        ["5", "Role-permission matrix"],
        ["6", "Database table dictionary"],
        ["7", "Representative action handlers"],
        ["8", "Security controls"],
        ["9", "Functional test cases"],
        ["10", "Boundary and negative tests"],
        ["11", "Deployment checklist"],
        ["12", "Risk register"],
    ])

    # 9 - Abbreviations
    new_report_page(doc, "Abbreviations and Terminology")
    table(doc, ["Term", "Meaning"], [
        ["CRUD", "Create, Read, Update and Delete"],
        ["CSRF", "Cross-Site Request Forgery"],
        ["DBMS", "Database Management System"],
        ["DFD", "Data Flow Diagram"],
        ["ERD", "Entity Relationship Diagram"],
        ["HTTP", "Hypertext Transfer Protocol"],
        ["MVC-like", "Lightweight separation of configuration, helpers, actions and views"],
        ["PDO", "PHP Data Objects"],
        ["RBAC", "Role-Based Access Control"],
        ["SQL", "Structured Query Language"],
        ["UI / UX", "User Interface / User Experience"],
        ["XAMPP", "Apache, MariaDB, PHP and related local development tools"],
    ])

    # 10 - Executive summary
    new_report_page(doc, "Executive Summary", "Project Synopsis")
    add_standard_narrative(
        doc, "The Cafe project",
        "the public website attracts visitors, the customer workspace supports purchase decisions, and operational dashboards move each accepted order through preparation and delivery",
        "all roles read and update one controlled application state",
        "a single demonstrable application that links customer experience to restaurant operations and persistent records",
    )
    bullets(doc, [
        "Problem addressed: fragmented and manual food-order handling.",
        "Primary users: customers, administrators, kitchen staff, delivery staff and cashiers.",
        "Core value: one shared database and controlled order-status lifecycle.",
        "Academic value: demonstrates frontend, backend, database, security, testing and deployment.",
    ])

    chapter_sections = [
        ("Chapter 1: Introduction", [
            ("1.1 Background of Online Food Ordering", "Digital ordering reduces telephone errors, presents an always-current menu, and provides status visibility to customers and staff.", "customers interact with structured menu, cart and checkout records", "orders can be processed consistently from placement through delivery"),
            ("1.2 Problem Statement", "manual ordering can produce pricing errors, missing instructions, duplicate communication and limited reporting.", "the application centralizes menu, customer and order information", "staff members operate from a shared source of truth"),
            ("1.3 Aim of the Project", "the aim is to design and implement a secure responsive food ordering and restaurant management application.", "the implementation combines customer ordering with role-based restaurant operations", "the complete food-delivery lifecycle is represented in one system"),
            ("1.4 Project Objectives", "objectives translate the broad aim into measurable functional and quality targets.", "each target is mapped to an implemented module and verification case", "project completion can be evaluated objectively"),
            ("1.5 Scope and Boundaries", "scope defines what the current academic implementation includes and what remains outside a local demonstration.", "real payment settlement, GPS telemetry and production email delivery are represented as integration points rather than live services", "the report remains accurate about implemented capability"),
            ("1.6 Significance of the Project", "the project demonstrates how practical restaurant processes can be expressed through data models and guarded web workflows.", "students can inspect the complete source, SQL, setup guide and role flows", "the system is useful for learning and portfolio demonstration"),
        ]),
        ("Chapter 2: Existing System and Feasibility", [
            ("2.1 Existing Manual System", "traditional ordering depends on calls, handwritten tickets and disconnected spreadsheets.", "staff repeat customer and item information across stages", "delays and transcription errors become difficult to trace"),
            ("2.2 Limitations of Existing Approaches", "static restaurant websites may display food but do not coordinate stock, coupons, order status or role permissions.", "Cafe adds database-backed operations and status transitions", "the website behaves as an operational application rather than a brochure"),
            ("2.3 Proposed System", "the proposed solution combines public discovery, secure customer actions and staff dashboards.", "a shared MariaDB schema records menu, users, orders, items, reviews and settings", "each role sees actions relevant to its responsibilities"),
            ("2.4 Technical Feasibility", "Core PHP, PDO, MariaDB, Bootstrap and XAMPP are mature and widely available.", "the project requires no Node.js build process or paid local service", "a beginner can install and run the application using the supplied guide"),
            ("2.5 Economic Feasibility", "the software stack is open source and can be demonstrated on ordinary personal computers.", "development and local deployment avoid license fees", "the project is economical for education and small proof-of-concept use"),
            ("2.6 Operational and Schedule Feasibility", "role-oriented interfaces match the sequence already familiar to restaurant staff.", "incremental modules can be demonstrated independently", "the project can be completed, tested and maintained within an academic schedule"),
        ]),
        ("Chapter 3: Requirements Engineering", [
            ("3.1 Stakeholder Identification", "stakeholders include customers, administrators, kitchen workers, delivery partners, cashiers and project maintainers.", "each stakeholder has a defined goal and data boundary", "requirements can be prioritized without mixing responsibilities"),
            ("3.2 Customer Functional Requirements", "customers require account access, menu discovery, cart control, coupon calculation, checkout, tracking, favorites and reviews.", "the storefront and customer workspace expose these functions through responsive pages", "the complete purchase journey can be demonstrated"),
            ("3.3 Administrative Requirements", "administrators require overview metrics, order control, menu CRUD, image upload, customer and staff management, coupons, moderation, reports and settings.", "dashboard routes read live database values and actions persist validated changes", "management features are not merely visual placeholders"),
            ("3.4 Kitchen and Delivery Requirements", "kitchen staff need a focused queue while delivery staff need assignment, contact, route and completion actions.", "role transitions restrict kitchen to preparation states and riders to delivery states", "operational responsibility remains clear"),
            ("3.5 Non-Functional Requirements", "quality requirements cover security, responsiveness, usability, maintainability, reliability, portability and performance.", "the implementation uses reusable helpers, shared styles, prepared statements and responsive breakpoints", "the application remains understandable and usable across devices"),
            ("3.6 Hardware and Software Requirements", "the local system requires a modern Windows computer, XAMPP, browser and adequate storage for uploads.", "the supplied SQL and default configuration match a standard XAMPP setup", "new users can reproduce the environment"),
        ]),
        ("Chapter 4: System Analysis and Design", [
            ("4.1 System Context", "the context boundary shows people and persistent services that exchange information with Cafe.", "customer and staff requests enter the application and validated records reach MariaDB", "external responsibilities are distinct from internal modules"),
            ("4.2 Layered Architecture", "the architecture separates browser presentation, PHP rendering, action handling, helpers, PDO and storage.", "GET requests render pages while POST requests pass through action and security checks", "the source remains lightweight without losing logical separation"),
            ("4.3 Customer Use Cases", "customer use cases begin with authentication and continue through menu, cart, checkout, tracking and review.", "directed relationships show the sequence of user intent", "functional coverage can be compared with implemented pages"),
            ("4.4 Staff Use Cases", "staff use cases are partitioned by administrative, kitchen and delivery duties.", "role guards and state transitions prevent unrelated actions", "least-privilege operation is supported"),
            ("4.5 Context-Level Data Flow", "the level-zero DFD summarizes how credentials, menu data, order data and status information move.", "processes communicate with explicit stores rather than hidden state", "data responsibilities can be reviewed before database design"),
            ("4.6 Detailed Order Data Flow", "the level-one order DFD expands validation, calculation, insertion, stock update, transaction and notification.", "every directed edge represents data or control passed to the next stage", "partial orders are avoided through transaction rollback"),
            ("4.7 Navigation Design", "navigation connects public discovery to customer pages and separates staff destinations by role.", "active sidebar states and protected redirects preserve context", "users reach permitted functionality with fewer ambiguous controls"),
            ("4.8 Order State Machine", "order status is modeled as a directional lifecycle rather than an arbitrary label.", "kitchen and delivery transitions have defined predecessors", "invalid jumps and cross-rider updates can be rejected"),
        ]),
        ("Chapter 5: Database Design", [
            ("5.1 Database Design Objectives", "the data model must retain identity, menu, financial, fulfillment and feedback records without unnecessary duplication.", "foreign keys link transactional rows and indexes support common status and category access", "records remain consistent and queryable"),
            ("5.2 Entity Relationship Model", "the ER model represents primary entities and foreign-key direction from child record to referenced parent.", "orders own line items while users, menu items and categories remain reusable entities", "deletion policies preserve historical meaning where appropriate"),
            ("5.3 User and Address Tables", "users store authentication identity, role, contact and account state while addresses store reusable delivery locations.", "password values contain hashes rather than plain text", "authentication and profile data remain centralized"),
            ("5.4 Menu and Category Tables", "categories group dishes and menu items store price, image, dietary, rating, preparation and stock properties.", "availability flags support soft removal from the public menu", "historical order items remain readable after menu changes"),
            ("5.5 Order and Order-Item Tables", "orders store customer snapshot, totals, payment, fulfillment and timestamps while order items preserve purchased name and price.", "transactional insertion and stock decrement occur together", "financial reconstruction and reporting are possible"),
            ("5.6 Supporting Tables", "coupons, reviews, favorites, notifications, helpful votes, settings and activity logs extend the core model.", "each supporting table has a narrow responsibility", "features can evolve without overloading order records"),
        ]),
        ("Chapter 6: Implementation", [
            ("6.1 Project Structure", "configuration, helpers, action handlers, customer pages, dashboards, assets, SQL, uploads and documentation are stored in recognizable locations.", "shared behavior is required from functions.php and configuration is centralized", "new contributors can locate responsibilities quickly"),
            ("6.2 Authentication Implementation", "login retrieves an active account by email and verifies a password hash before regenerating the session ID.", "the session stores role identity used by guards and navigation", "credentials are never compared as plain text"),
            ("6.3 Customer Interface Implementation", "the public landing page, menu, item details, cart, checkout, dashboard, favorites, orders and tracking pages are rendered from live data.", "forms map to named actions and visible totals are recalculated on both client and server", "customer actions remain functional as well as attractive"),
            ("6.4 Cart, Coupon and Checkout Logic", "cart state is held in the session while product prices, availability and stock are revalidated before order creation.", "coupon rules are recalculated against the current subtotal and checkout enforces minimum value", "tampered or stale cart information cannot silently create invalid orders"),
            ("6.5 Administrative Dashboard", "administrators receive real totals, paginated menu records, CSV export and CRUD interfaces.", "image uploads use MIME checks, random filenames and size limits", "dashboard controls persist changes instead of acting as placeholders"),
            ("6.6 Kitchen Dashboard", "confirmed orders enter the kitchen queue where staff can begin preparation and mark completion.", "the interface displays item quantity, instructions and elapsed time", "the kitchen is isolated from unrelated administrative functions"),
            ("6.7 Delivery Dashboard", "ready orders are offered to delivery staff and assignment occurs when delivery begins.", "another rider cannot update an already assigned delivery", "the final action records delivery and settles cash-on-delivery status"),
            ("6.8 Reporting and Settings", "reports aggregate revenue, order values and repeat activity while settings control contact, hours, fees, thresholds and tax.", "CSV export uses filtered order data and settings are upserted through a whitelist", "operational data can be reviewed and adjusted safely"),
        ]),
        ("Chapter 7: Security and Reliability", [
            ("7.1 Authentication and Authorization", "authentication establishes identity and authorization limits every privileged action by role.", "customer-only routes redirect staff to dashboards and inactive accounts cannot sign in", "data and functions are protected beyond visual menu hiding"),
            ("7.2 CSRF and Session Security", "every state-changing form includes a session token checked before action dispatch.", "empty or mismatched tokens produce a 419 response and successful login regenerates the session identifier", "forged cross-site form submissions are rejected"),
            ("7.3 SQL Injection and XSS Prevention", "PDO prepared statements separate SQL structure from values and output uses HTML escaping.", "dynamic redirects are restricted to recognized local prefixes", "common injection and stored markup risks are reduced"),
            ("7.4 Transaction and Stock Integrity", "order creation validates current menu values and updates order, line items, stock, coupon count and notification in one transaction.", "a failed stock update throws an exception and rolls back", "the database avoids partially completed orders"),
            ("7.5 Upload and Validation Controls", "uploaded images are restricted by MIME type, size and randomized server filename.", "numeric, date, email, role and text-length checks are applied before storage", "invalid content fails with an understandable response"),
        ]),
        ("Chapter 8: Testing and Validation", [
            ("8.1 Testing Strategy", "testing combines PHP syntax checks, JavaScript parsing, HTTP flows, database queries, role sign-in tests and cleanup.", "each test follows the browser-to-handler-to-database-to-response story", "evidence covers more than page rendering"),
            ("8.2 Authentication Tests", "valid and invalid credentials, inactive users, customer redirects and staff dashboards are checked.", "CSRF omission and unauthorized export access are tested separately", "identity boundaries behave as expected"),
            ("8.3 Cart and Checkout Tests", "tests add several items, change quantity, verify live markup, apply and remove coupons, reach checkout and place an order.", "database records and stock changes are inspected before temporary test data is removed", "calculation and persistence agree"),
            ("8.4 Multi-Role Lifecycle Test", "a test order is confirmed by an administrator, prepared by kitchen staff, accepted by a rider and completed.", "the final database record contains delivered status, assigned rider and paid COD state", "the principal business workflow is verified end to end"),
            ("8.5 CRUD and Upload Tests", "temporary menu, staff and coupon records exercise create, update, status toggle and delete paths.", "a real PNG is submitted through multipart upload and served back through the web server", "management forms and storage behavior are confirmed"),
            ("8.6 Review, Settings and Export Tests", "customer review submission, admin reply, helpful vote, visibility moderation, settings persistence and CSV response are checked.", "content type and downloaded byte length provide export evidence", "secondary features are verified rather than assumed"),
            ("8.7 Responsive and Usability Review", "desktop and mobile breakpoints are reviewed for navigation, cards, forms, tables and persistent sidebars.", "cache-busting asset versions prevent stale CSS from hiding visual fixes", "the current interface remains usable across common viewport sizes"),
            ("8.8 Test Result Summary", "all critical flows pass after fixes for headers, permissions, credential exposure, dropdown overlap and calculated totals.", "the database is restored to clean demo data after destructive tests", "the repository remains reproducible for viewers and examiners"),
        ]),
        ("Chapter 9: Deployment and User Manual", [
            ("9.1 Deployment Architecture", "the local deployment uses a browser, Apache, PHP runtime, project folder and MariaDB service.", "phpMyAdmin imports the supplied schema and uploaded images are stored under the project", "the environment can be reproduced without a commercial service"),
            ("9.2 Installation Procedure", "installation requires XAMPP, copying the project to htdocs, starting Apache and MySQL, importing SQL and opening localhost.", "the README provides exact paths, URLs and troubleshooting advice", "non-technical users can follow a linear setup process"),
            ("9.3 Customer User Manual", "customers register or sign in, browse categories, search, add items, modify quantities, apply coupons, checkout and track.", "delivered orders expose review forms and favorites remain available from the sidebar", "the user journey is consistent from desktop to mobile"),
            ("9.4 Administrator User Manual", "administrators sign in to a role dashboard and use menu, orders, customers, team, coupons, reviews, reports and settings pages.", "edit and archive actions use confirmation and all state-changing requests carry CSRF tokens", "operational tasks are grouped by navigation item"),
            ("9.5 Kitchen and Delivery Manuals", "kitchen users process confirmed tickets while delivery users work only with ready or assigned orders.", "buttons advance the allowed next state and customer notifications record updates", "specialized roles require minimal training"),
            ("9.6 Backup and Maintenance", "database exports, upload backups, credential changes and periodic review of dependencies are recommended.", "production deployment must replace default root credentials and demo accounts", "operational continuity and privacy are improved"),
        ]),
        ("Chapter 10: Results, Limitations and Future Work", [
            ("10.1 Results and Achievements", "the final system integrates attractive public, customer and staff experiences with functional server-side behavior.", "fifteen database tables and multiple roles support the full order lifecycle", "project objectives are substantially achieved"),
            ("10.2 Current Limitations", "online payment options are demonstrations, delivery location is not real-time GPS, and outbound email is not configured.", "the project targets one restaurant and local deployment", "limitations are explicit rather than disguised as completed integrations"),
            ("10.3 Future Enhancements", "future work can add Razorpay or Stripe, maps, delivery OTP, email/SMS, WebSocket updates, multiple branches, taxes by jurisdiction and automated tests.", "the existing schema and role architecture provide extension points", "the academic prototype can evolve into a hosted service"),
            ("10.4 Conclusion", "Cafe demonstrates the design, implementation and verification of a practical database-backed food ordering system.", "customer convenience is connected to controlled restaurant workflows and reporting", "the project provides a strong foundation for further full-stack study"),
        ]),
    ]

    diagram_by_title = {item["title"]: item for item in diagrams}
    diagram_order = [
        "System Context Diagram", "Layered Application Architecture",
        "Customer Use-Case Diagram", "Role-Based Staff Use Cases",
        "Data Flow Diagram - Level 0", "Order Processing DFD - Level 1",
        "Entity Relationship Diagram", "Login Sequence",
        "Checkout and Order Creation Sequence", "Order Status State Machine",
        "Local Deployment Diagram", "Module Dependency Diagram",
        "Request Security Validation Flow", "Application Navigation Structure",
        "Verification and Test Pipeline",
    ]
    section_to_diagram = {
        "4.1 System Context": diagram_order[0],
        "4.2 Layered Architecture": diagram_order[1],
        "4.3 Customer Use Cases": diagram_order[2],
        "4.4 Staff Use Cases": diagram_order[3],
        "4.5 Context-Level Data Flow": diagram_order[4],
        "4.6 Detailed Order Data Flow": diagram_order[5],
        "4.7 Navigation Design": diagram_order[13],
        "4.8 Order State Machine": diagram_order[9],
        "5.2 Entity Relationship Model": diagram_order[6],
        "6.1 Project Structure": diagram_order[11],
        "6.2 Authentication Implementation": diagram_order[7],
        "6.4 Cart, Coupon and Checkout Logic": diagram_order[8],
        "7.2 CSRF and Session Security": diagram_order[12],
        "8.1 Testing Strategy": diagram_order[14],
        "9.1 Deployment Architecture": diagram_order[10],
    }

    figure_number = 1
    for chapter, sections in chapter_sections:
        for heading, details, operation, outcome in sections:
            new_report_page(doc, heading, chapter)
            add_standard_narrative(doc, heading.split(":")[-1].strip(), details, operation, outcome)
            if heading in section_to_diagram:
                diag = diagram_by_title[section_to_diagram[heading]]
                figure(doc, ROOT / diag["file"], diag["title"], figure_number)
                paragraph(doc, "Arrow validation note: arrowheads point from the initiating or child component toward the receiving, referenced, or next-state component. Every endpoint was programmatically checked before the figure was embedded.")
                figure_number += 1
            if heading == "1.4 Project Objectives":
                bullets(doc, [
                    "Implement secure registration, login, logout and role guards.",
                    "Provide searchable and categorized food discovery.",
                    "Calculate cart, discounts, delivery fee, tax and final total.",
                    "Persist transactional orders and enforce stock availability.",
                    "Provide separate admin, kitchen, delivery and staff experiences.",
                    "Support reporting, exports, reviews, settings and responsive design.",
                ], numbered=True)
            elif heading == "3.5 Non-Functional Requirements":
                table(doc, ["Quality", "Requirement", "Implementation Evidence"], [
                    ["Security", "Reject forged and unauthorized changes", "CSRF, role guards, PDO, escaping"],
                    ["Usability", "Clear actions on desktop and mobile", "Responsive navigation and forms"],
                    ["Reliability", "Avoid partial order writes", "PDO transaction and rollback"],
                    ["Maintainability", "Central shared helpers", "config.php and functions.php"],
                    ["Portability", "Run on common local stack", "XAMPP README and SQL seed"],
                ])
            elif heading == "3.6 Hardware and Software Requirements":
                table(doc, ["Category", "Minimum", "Recommended"], [
                    ["Processor", "Dual-core 2 GHz", "Modern quad-core CPU"],
                    ["Memory", "4 GB RAM", "8 GB RAM"],
                    ["Storage", "500 MB free", "2 GB including uploads"],
                    ["Operating system", "Windows 10", "Windows 11"],
                    ["Server stack", "PHP 8.1 + MariaDB", "Current XAMPP"],
                    ["Browser", "Modern Chromium/Firefox", "Latest Chrome or Edge"],
                ])
            elif heading == "3.3 Administrative Requirements":
                table(doc, ["Module", "Create", "Read", "Update", "Archive/Delete"], [
                    ["Menu", "Yes", "Yes", "Yes", "Archive"],
                    ["Staff", "Yes", "Yes", "Yes", "Deactivate"],
                    ["Coupons", "Yes", "Yes", "Yes", "Delete"],
                    ["Reviews", "Reply", "Yes", "Visibility", "Hide"],
                    ["Orders", "System", "Yes", "Status", "Cancel"],
                    ["Settings", "Upsert", "Yes", "Yes", "Not applicable"],
                ])
            elif heading == "5.6 Supporting Tables":
                table(doc, ["Table", "Purpose", "Important Fields"], [
                    ["coupons", "Discount rules", "code, type, value, minimum_order, valid_until"],
                    ["reviews", "Verified feedback", "rating, comment, admin_reply, helpful_count"],
                    ["favorites", "Saved dishes", "user_id, menu_item_id"],
                    ["notifications", "Customer updates", "title, message, icon, is_read"],
                    ["settings", "Restaurant configuration", "setting_key, setting_value"],
                    ["activity_logs", "Privileged audit support", "user_id, action, ip_address"],
                ])
            elif heading == "6.1 Project Structure":
                table(doc, ["Path", "Responsibility"], [
                    ["index.php", "Public and customer routes"],
                    ["dashboard.php", "Role-specific dashboards"],
                    ["actions.php", "POST action controller"],
                    ["functions.php", "Shared auth, CSRF, cart and query helpers"],
                    ["config.php", "Session and PDO configuration"],
                    ["database/food_ordering.sql", "Schema and demo data"],
                    ["assets/", "Styles, scripts and visual assets"],
                    ["uploads/menu/", "Validated menu uploads"],
                ])
            elif heading == "7.5 Upload and Validation Controls":
                table(doc, ["Control", "Rule"], [
                    ["Image MIME", "JPEG, PNG, WebP or GIF"],
                    ["Maximum image size", "2 MB application limit"],
                    ["Filename", "Random 24-character hexadecimal value"],
                    ["Coupon code", "Uppercase letters, digits, underscore or hyphen"],
                    ["Staff role", "Admin, kitchen, delivery or staff only"],
                    ["Review comment", "5 to 500 characters"],
                ])

    # Test case pages
    new_report_page(doc, "Detailed Functional Test Cases", "Chapter 8: Testing and Validation")
    test_rows = [
        ["TC-01", "Customer login", "Valid seeded credentials", "Customer home/dashboard", "Pass"],
        ["TC-02", "Invalid login", "Wrong password", "Error flash; no session", "Pass"],
        ["TC-03", "Add cart item", "Available menu ID", "Session quantity created", "Pass"],
        ["TC-04", "Update quantity", "Quantity 3", "Line and grand totals update", "Pass"],
        ["TC-05", "Apply coupon", "WELCOME100 over minimum", "Discount applied", "Pass"],
        ["TC-06", "Remove coupon", "Active coupon", "Discount removed", "Pass"],
        ["TC-07", "Place order", "Valid cart/address", "Order and items committed", "Pass"],
        ["TC-08", "Kitchen transition", "Confirmed order", "Preparing then ready", "Pass"],
        ["TC-09", "Delivery transition", "Ready order", "Rider assigned and delivered", "Pass"],
        ["TC-10", "Menu image upload", "Valid PNG under 2 MB", "Random local image path", "Pass"],
        ["TC-11", "CSV export", "Admin authenticated", "CSV content-type and rows", "Pass"],
        ["TC-12", "Missing CSRF", "POST without token", "HTTP 419", "Pass"],
        ["TC-13", "Staff edit", "Existing kitchen user", "Updated persistent profile", "Pass"],
        ["TC-14", "Review reply", "Admin response", "Reply visible to customer", "Pass"],
        ["TC-15", "Settings save", "Valid fee/tax values", "Upserted settings", "Pass"],
    ]
    table(doc, ["ID", "Scenario", "Input", "Expected Result", "Status"], test_rows)

    new_report_page(doc, "Boundary and Negative Test Cases", "Chapter 8: Testing and Validation")
    table(doc, ["ID", "Boundary / Invalid Input", "Expected Protection", "Result"], [
        ["NT-01", "Empty CSRF token", "Reject with 419", "Pass"],
        ["NT-02", "Quantity above 10", "Clamp to 10", "Pass"],
        ["NT-03", "Quantity exceeds stock", "Reject checkout", "Pass"],
        ["NT-04", "Expired coupon", "Reject or remove", "Pass"],
        ["NT-05", "Order below minimum", "Reject order placement", "Pass"],
        ["NT-06", "Non-image upload", "Reject MIME type", "Pass"],
        ["NT-07", "Image above size limit", "Reject upload", "Pass"],
        ["NT-08", "Kitchen attempts delivered", "Reject transition", "Pass"],
        ["NT-09", "Second rider updates assignment", "Reject ownership", "Pass"],
        ["NT-10", "Duplicate review", "Reject second submission", "Pass"],
        ["NT-11", "Unauthenticated export", "Redirect to login", "Pass"],
        ["NT-12", "Unknown action", "Safe error flash", "Pass"],
    ])

    # Risk page
    new_report_page(doc, "Project Risk Register", "Project Management")
    table(doc, ["Risk", "Probability", "Impact", "Mitigation"], [
        ["Database service not running", "Medium", "High", "Clear XAMPP checks and error page"],
        ["Port 80 conflict", "Medium", "Medium", "README port 8080 procedure"],
        ["Stale browser assets", "Medium", "Low", "filemtime cache versioning"],
        ["Invalid uploads", "Medium", "High", "MIME, size and randomized filename"],
        ["Unauthorized status change", "Low", "High", "Role and state-transition guards"],
        ["Stock race during checkout", "Low", "High", "Conditional update in transaction"],
        ["Demo credentials used publicly", "Medium", "High", "Change credentials before production"],
        ["External image outage", "Medium", "Low", "Local upload and placeholder support"],
    ])

    # Bibliography
    new_report_page(doc, "Bibliography", "References")
    references = [
        "PHP Documentation. PHP Manual: Sessions, password_hash, PDO and file uploads. https://www.php.net/docs.php",
        "MariaDB Documentation. SQL syntax, transactions, foreign keys and indexes. https://mariadb.com/kb/en/documentation/",
        "Apache Friends. XAMPP installation and local server environment. https://www.apachefriends.org/",
        "Bootstrap Documentation. Responsive grid, components and utilities. https://getbootstrap.com/docs/5.3/",
        "Bootstrap Icons. Open-source icon library. https://icons.getbootstrap.com/",
        "OWASP Foundation. Cross-Site Request Forgery Prevention Cheat Sheet.",
        "OWASP Foundation. SQL Injection Prevention Cheat Sheet.",
        "OWASP Foundation. Cross Site Scripting Prevention Cheat Sheet.",
        "Ian Sommerville. Software Engineering. Requirements, architecture and testing principles.",
        "Ramez Elmasri and Shamkant Navathe. Fundamentals of Database Systems.",
        "Roger S. Pressman. Software Engineering: A Practitioner's Approach.",
        "MDN Web Docs. HTML forms, responsive design and JavaScript APIs. https://developer.mozilla.org/",
    ]
    bullets(doc, references, numbered=True)

    # Appendix A
    new_report_page(doc, "Appendix A: Database Table Dictionary", "Appendices")
    table(doc, ["Table", "Primary Key", "Main Relationships", "Role"], [
        ["users", "id", "Parent of addresses, orders, reviews, favorites", "Identity and RBAC"],
        ["addresses", "id", "user_id -> users", "Saved delivery location"],
        ["categories", "id", "Parent of menu_items", "Menu organization"],
        ["menu_items", "id", "category_id -> categories", "Dish catalogue and stock"],
        ["coupons", "id", "Referenced by code snapshot", "Discount rules"],
        ["orders", "id", "user_id and delivery_user_id -> users", "Order header"],
        ["order_items", "id", "order_id -> orders; menu_item_id -> menu_items", "Purchased lines"],
        ["reviews", "id", "user, order and menu foreign keys", "Verified feedback"],
        ["favorites", "Composite", "user and menu foreign keys", "Saved dishes"],
        ["notifications", "id", "user_id -> users", "Status updates"],
        ["activity_logs", "id", "user_id -> users", "Audit support"],
        ["settings", "setting_key", "Independent", "Restaurant configuration"],
    ])

    # Appendix B
    new_report_page(doc, "Appendix B: Representative Action Catalogue", "Appendices")
    table(doc, ["Action", "Authorized Role", "Persistent Effect"], [
        ["login", "Any active account", "Session identity and activity log"],
        ["register", "Public", "New customer user"],
        ["add_cart / update_cart", "Public session", "Session cart"],
        ["apply_coupon / remove_coupon", "Public session", "Session discount"],
        ["place_order", "Customer", "Order, lines, stock, coupon count, notification"],
        ["submit_review", "Customer", "Review and menu rating"],
        ["order_status", "Admin / staff roles", "Status, rider and payment state"],
        ["save_item", "Admin", "Menu insert/update and optional upload"],
        ["save_staff / toggle_staff", "Admin", "Staff account state"],
        ["save_coupon / toggle_coupon", "Admin", "Coupon rules"],
        ["reply_review / toggle_review", "Admin", "Moderation state"],
        ["save_settings", "Admin", "Whitelisted settings upsert"],
    ])

    # Appendix C
    new_report_page(doc, "Appendix C: Installation Checklist", "Appendices")
    bullets(doc, [
        "Install XAMPP in C:\\xampp with Apache, PHP, MySQL/MariaDB and phpMyAdmin.",
        "Copy the project to C:\\xampp\\htdocs\\food-ordering-website.",
        "Start Apache and MySQL from the XAMPP Control Panel.",
        "Open http://localhost/phpmyadmin and import database/food_ordering.sql.",
        "Open http://localhost/food-ordering-website/.",
        "Use README.md for demo accounts and troubleshooting.",
        "Change all credentials before any public deployment.",
    ], numbered=True)
    paragraph(doc, "Installation acceptance criteria: the public home returns HTTP 200, the menu displays seeded items, customer and staff accounts authenticate, and the database contains all expected tables.")

    # Appendix D
    new_report_page(doc, "Appendix D: Diagram Arrow Validation", "Appendices")
    table(doc, ["Diagram", "Nodes", "Directed Arrows", "Validation"], [
        [d["title"], d["node_count"], d["arrow_count"], "PASS"] for d in diagrams
    ])
    paragraph(doc, "Validation algorithm: each edge is stored as (source, target, label). Generation fails when the source or target is absent, when a self-loop is found, or when an edge label is empty. NetworkX draws an -|> arrowhead toward the target node. Therefore, the visible arrow direction matches the declared semantic direction.")

    # Appendix E
    new_report_page(doc, "Appendix E: Submission Customization Checklist", "Appendices")
    bullets(doc, [
        "Replace institution, department, roll number and supervisor placeholders.",
        "Insert official college logo on the title page if required.",
        "Collect signatures on certificate and declaration pages.",
        "Update the Word table of contents and verify final pagination.",
        "Replace or add screenshots required by the department template.",
        "Review citation style against institutional guidelines.",
        "Export a final PDF and inspect every diagram at 100% zoom.",
        "Confirm repository and source-code links remain accessible.",
    ])

    # Final page
    new_report_page(doc, "End of Report")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run(PROJECT)
    r.bold = True
    r.font.size = Pt(18)
    paragraph(doc, "This report was generated from the implemented Cafe project and includes programmatically validated directed diagrams. Institution-specific placeholders remain editable in Microsoft Word.")

    doc.core_properties.title = PROJECT
    doc.core_properties.subject = "College full-stack PHP and MySQL project report"
    doc.core_properties.author = STUDENT
    doc.core_properties.keywords = "PHP, MySQL, food ordering, restaurant management, project report"
    doc.core_properties.comments = "Generated report with validated directed diagrams."
    doc.save(OUTPUT)


def main():
    validations = []
    for spec in diagram_specs():
        validations.append(validate_and_render(spec))
    VALIDATION.write_text(json.dumps({
        "status": "PASS",
        "diagram_count": len(validations),
        "total_nodes": sum(item["node_count"] for item in validations),
        "total_directed_arrows": sum(item["arrow_count"] for item in validations),
        "diagrams": validations,
    }, indent=2), encoding="utf-8")
    build_document(validations)
    print(f"Created: {OUTPUT}")
    print(f"Diagrams validated: {len(validations)}")
    print(f"Directed arrows validated: {sum(item['arrow_count'] for item in validations)}")


if __name__ == "__main__":
    main()
